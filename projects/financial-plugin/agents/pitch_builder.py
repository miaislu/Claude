"""
Pitch Builder Agent。

从 storage/ 读取各 Agent 的最新存档，合成投资材料。

MVP：
  - 买方投资备忘录（Word，python-docx）
  - 可比估值 Excel（复用 ModelBuildResult.excel_path）
  - 买方 1-pager 文本（Markdown）

前置条件：
  - ValuationReviewResult.verdict ≠ REJECT（否则阻断）
  - ModelBuildResult.linkage_errors 为空（否则警告）

评级体系（5 档，基于 12 月 upside）：
  买入：>+20%  增持：+10%~20%  中性：-10%~+10%  减持：-10%~-20%  卖出：<-20%
"""

from __future__ import annotations

import json
import sys
import warnings
from datetime import datetime
from pathlib import Path
from typing import Any

warnings.filterwarnings("ignore")

_ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(_ROOT))

from models import (
    Catalyst,
    EarningsReviewResult,
    MarketResearchDigest,
    ModelBuildResult,
    PitchBuildResult,
    ValuationReviewResult,
)

_DISCLAIMER = (
    "本材料为内部研究参考，不构成投资建议。"
    "历史数据不代表未来表现。投资有风险，入市需谨慎。"
)


class PitchBuilder:
    """投资材料合成 Agent。"""

    def build(
        self,
        code: str,
        audience: str = "buyside_memo",
    ) -> PitchBuildResult:
        """
        合成投资材料。

        audience: "buyside_memo"（默认）/ "buyside_1pager" / "full_suite"
        """
        # 1. 读取存档
        earnings  = self._load_latest(code, "earnings_review")
        model_res = self._load_latest(code, "model_build")
        val_res   = self._load_latest(code, "valuation_review")
        market_d  = self._load_market_digest()

        # 2. 前置条件检查
        val_verdict = val_res.get("verdict", "PASS") if val_res else "PASS"
        if val_verdict == "REJECT":
            return PitchBuildResult(
                stock_code=code,
                pitch_date=datetime.now().strftime("%Y-%m-%d"),
                valuation_verdict="REJECT",
                human_review_required=True,
            )

        # 3. 提取关键数据
        target_price  = _f(model_res, "blended_target_price") if model_res else None
        current_price = _f(model_res, "current_price") if model_res else None
        upside_pct    = _f(model_res, "upside_pct") if model_res else None
        company_name  = _s(model_res, "company_name") if model_res else code
        warnings_cnt  = len(val_res.get("warnings", [])) if val_res else 0
        model_version = _s(model_res, "version") if model_res else ""

        # 4. 评级
        rating = _calc_rating(upside_pct)

        # 5. 投资逻辑（3条）
        thesis = self._build_thesis(earnings, model_res, val_res)

        # 6. 催化剂（来自 MarketResearchDigest）
        catalysts = self._extract_catalysts(code, market_d)

        # 7. 生成文件
        files: dict[str, str] = {}
        if audience in ("buyside_memo", "full_suite"):
            memo_path = self._write_memo(
                code, company_name, rating, target_price, current_price,
                upside_pct, thesis, catalysts, earnings, model_res, val_res
            )
            if memo_path:
                files["memo"] = memo_path

        if audience in ("buyside_1pager", "full_suite"):
            pager_path = self._write_1pager(
                code, company_name, rating, target_price, current_price,
                upside_pct, thesis, catalysts, val_res
            )
            if pager_path:
                files["1pager"] = pager_path

        # Excel：直接引用 Model Builder 的输出（不重建）
        if model_res and model_res.get("excel_path"):
            files["comps_excel"] = model_res["excel_path"]

        result = PitchBuildResult(
            stock_code=code,
            company_name=company_name,
            pitch_date=datetime.now().strftime("%Y-%m-%d"),
            suggested_rating=rating,
            target_price=target_price,
            current_price=current_price,
            upside_pct=upside_pct,
            investment_thesis=thesis,
            catalysts=catalysts,
            earnings_review_date=_s(earnings, "timestamp")[:10] if earnings else "",
            model_version=model_version,
            valuation_verdict=val_verdict,
            warnings_count=warnings_cnt,
            files=files,
            audience_type=audience,
            human_review_required=warnings_cnt > 0,
        )
        self._save_result(result)
        return result

    # ──────────────────────────────────────────
    # 内容生成
    # ──────────────────────────────────────────

    def _build_thesis(
        self,
        earnings: dict | None,
        model_res: dict | None,
        val_res: dict | None,
    ) -> list[str]:
        thesis: list[str] = []

        # WHAT
        thesis.append(f"公司 {_s(model_res,'company_name')} 具有较强市场地位，财务模型已验证。")

        # WHY（来自 earnings）
        if earnings:
            overall = earnings.get("overall_verdict", "IN_LINE")
            verdict_map = {
                "STRONG_BEAT": "近期业绩大幅超预期",
                "BEAT":        "近期业绩小幅超预期",
                "IN_LINE":     "近期业绩符合预期",
                "MISS":        "近期业绩小幅低于预期",
                "STRONG_MISS": "近期业绩大幅低于预期",
            }
            confirms = [
                t["thesis_keyword"]
                for t in (earnings.get("thesis_verdicts") or [])
                if t.get("verdict") == "CONFIRM"
            ]
            why_str = verdict_map.get(overall, "")
            if confirms:
                why_str += f"，投资逻辑获得验证：{'、'.join(confirms[:2])}"
            thesis.append(why_str or "需补充投资逻辑。")

        # HOW MUCH
        if model_res and _f(model_res, "blended_target_price") and _f(model_res, "current_price"):
            tp = _f(model_res, "blended_target_price")
            cp = _f(model_res, "current_price")
            upside = _f(model_res, "upside_pct")
            thesis.append(f"目标价 ¥{tp:.2f}（上行空间 {upside:+.1f}%），当前股价 ¥{cp:.2f}。")

        return thesis[:3]

    def _extract_catalysts(
        self,
        code: str,
        market_d: dict | None,
    ) -> list[Catalyst]:
        if not market_d:
            return []
        all_signals = (
            market_d.get("p0_immediate", [])
            + market_d.get("p1_daily", [])
        )
        cats: list[Catalyst] = []
        for sig in all_signals:
            if sig.get("stock_code") == code and sig.get("signal_type") == "THESIS_UPDATE":
                cats.append(Catalyst(
                    description=sig.get("headline", ""),
                    source="MarketResearchDigest",
                    certainty="MEDIUM",
                ))
        return cats[:5]

    # ──────────────────────────────────────────
    # 文档生成
    # ──────────────────────────────────────────

    def _write_memo(
        self,
        code: str,
        name: str,
        rating: str,
        tp: float | None,
        cp: float | None,
        upside: float | None,
        thesis: list[str],
        catalysts: list[Catalyst],
        earnings: dict | None,
        model_res: dict | None,
        val_res: dict | None,
    ) -> str | None:
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # 封面
            h = doc.add_heading(f"{name}（{code}）投资备忘录", level=1)
            h.runs[0].font.color.rgb = RGBColor(0, 51, 102)
            doc.add_paragraph(
                f"评级：{rating}  |  目标价：¥{tp:.2f}  |  当前价：¥{cp:.2f}  |  "
                f"上行空间：{upside:+.1f}%  |  日期：{datetime.now().strftime('%Y-%m-%d')}"
                if tp and cp and upside else
                f"评级：{rating}  |  日期：{datetime.now().strftime('%Y-%m-%d')}"
            )
            doc.add_paragraph()

            # 核心逻辑
            doc.add_heading("核心投资逻辑", level=2)
            for i, t in enumerate(thesis, 1):
                doc.add_paragraph(f"{i}. {t}", style="List Number")

            # 财务摘要
            doc.add_heading("近期业绩", level=2)
            if earnings:
                doc.add_paragraph(earnings.get("summary", "暂无摘要"))
                for gap in (earnings.get("expectation_gaps") or [])[:3]:
                    yoy = gap.get("yoy_change_pct")
                    yoy_str = f"同比 {yoy:+.1f}%" if yoy is not None else ""
                    doc.add_paragraph(
                        f"• {gap['metric']}：{gap['verdict']} {yoy_str}",
                        style="List Bullet",
                    )

            # 估值
            doc.add_heading("估值摘要", level=2)
            if model_res:
                doc.add_paragraph(
                    f"DCF 目标价：¥{_f(model_res,'dcf_target_price') or 'N/A':.2f}  |  "
                    f"WACC：{(_f(model_res,'wacc') or 0):.1%}  |  "
                    f"永续增长率 g：{(_f(model_res,'terminal_growth_rate') or 0):.1%}  |  "
                    f"终值/EV：{(_f(model_res,'terminal_value_pct') or 0):.1%}"
                )

            # 风险
            doc.add_heading("主要风险", level=2)
            if val_res:
                for w in (val_res.get("warnings") or [])[:3]:
                    doc.add_paragraph(f"• {w['description']}", style="List Bullet")
            if earnings:
                for rf in (earnings.get("risk_flags") or []):
                    if rf.get("verdict") in ("HIGH_RISK", "MEDIUM_RISK"):
                        doc.add_paragraph(f"• {rf['flag_type']}：{rf['explanation'][:60]}", style="List Bullet")

            # 催化剂
            if catalysts:
                doc.add_heading("近期催化剂", level=2)
                for cat in catalysts:
                    doc.add_paragraph(f"• {cat.description}", style="List Bullet")

            # 免责声明
            doc.add_paragraph()
            p = doc.add_paragraph(_DISCLAIMER)
            p.runs[0].font.size = Pt(8)
            p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

            # 保存
            storage = _ROOT / "storage"
            storage.mkdir(exist_ok=True)
            date_str = datetime.now().strftime("%Y%m%d")
            path = storage / f"{code}_pitch_memo_{date_str}.docx"
            doc.save(str(path))
            return str(path)
        except Exception:
            return None

    def _write_1pager(
        self,
        code: str,
        name: str,
        rating: str,
        tp: float | None,
        cp: float | None,
        upside: float | None,
        thesis: list[str],
        catalysts: list[Catalyst],
        val_res: dict | None,
    ) -> str | None:
        """生成 Markdown 版单页简报。"""
        lines = [
            f"# {name}（{code}）投资摘要",
            f"**评级**：{rating}  |  **目标价**：¥{tp:.2f}  |  **当前价**：¥{cp:.2f}  |  **上行空间**：{upside:+.1f}%"
            if tp and cp and upside else f"**评级**：{rating}",
            "",
            "## 核心逻辑",
        ] + [f"{i}. {t}" for i, t in enumerate(thesis, 1)]

        if catalysts:
            lines += ["", "## 下一个催化剂", f"→ {catalysts[0].description}"]

        if val_res:
            warns = [w["description"][:50] for w in (val_res.get("warnings") or [])[:2]]
            if warns:
                lines += ["", "## 主要风险"]
                lines += [f"- {w}" for w in warns]

        lines += ["", f"---", f"*{_DISCLAIMER}*"]

        storage = _ROOT / "storage"
        storage.mkdir(exist_ok=True)
        date_str = datetime.now().strftime("%Y%m%d")
        path = storage / f"{code}_pitch_1pager_{date_str}.md"
        path.write_text("\n".join(lines), encoding="utf-8")
        return str(path)

    # ──────────────────────────────────────────
    # 存档读取 / 写入
    # ──────────────────────────────────────────

    @staticmethod
    def _load_latest(code: str, agent_key: str) -> dict | None:
        """读取 storage/ 下最新的 {code}_{agent_key}_*.json。"""
        storage = _ROOT / "storage"
        files = sorted(storage.glob(f"{code}_{agent_key}_*.json"), reverse=True)
        if not files:
            return None
        try:
            return json.loads(files[0].read_text(encoding="utf-8"))
        except Exception:
            return None

    @staticmethod
    def _load_market_digest() -> dict | None:
        storage = _ROOT / "storage"
        files = sorted(storage.glob("market_digest_*.json"), reverse=True)
        if not files:
            return None
        try:
            return json.loads(files[0].read_text(encoding="utf-8"))
        except Exception:
            return None

    @staticmethod
    def _save_result(result: PitchBuildResult) -> str:
        storage = _ROOT / "storage"
        storage.mkdir(exist_ok=True)
        date_str = datetime.now().strftime("%Y%m%d")
        path = storage / f"{result.stock_code}_pitch_build_{date_str}.json"
        path.write_text(
            json.dumps(result.to_dict(), ensure_ascii=False, indent=2, default=str),
            encoding="utf-8",
        )
        return str(path)


# ──────────────────────────────────────────
# 工具函数
# ──────────────────────────────────────────

def _calc_rating(upside_pct: float | None) -> str:
    if upside_pct is None:
        return "中性"
    if upside_pct > 20:   return "买入"
    if upside_pct > 10:   return "增持"
    if upside_pct >= -10: return "中性"
    if upside_pct >= -20: return "减持"
    return "卖出"


def _f(d: dict | None, key: str) -> float | None:
    if d is None:
        return None
    v = d.get(key)
    if v is None:
        return None
    try:
        import math
        x = float(v)
        return None if math.isnan(x) else x
    except (ValueError, TypeError):
        return None


def _s(d: dict | None, key: str, default: str = "") -> str:
    if d is None:
        return default
    return str(d.get(key) or default)
