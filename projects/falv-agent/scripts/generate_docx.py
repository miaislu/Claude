#!/usr/bin/env python3
"""
中国法律 AI Agent — Word 报告生成器（法律文书样式）
将 Markdown 格式的审查意见转换为正式法律文书风格的 Word 文档

依赖：pip3 install python-docx
用法：python3 generate_docx.py --input report.md --name "项目名称" --output /path/to/reports/
"""

import argparse
import os
import re
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm


# ── 配色：正式法律文书风格，无彩色 ────────────────────────────────────────
BLACK       = RGBColor(0x00, 0x00, 0x00)   # 正文
DARK_GRAY   = RGBColor(0x33, 0x33, 0x33)   # 标题
MED_GRAY    = RGBColor(0x66, 0x66, 0x66)   # 次级信息、注释
LIGHT_GRAY  = RGBColor(0xAA, 0xAA, 0xAA)  # 辅助线
HEADER_LINE = RGBColor(0x1A, 0x1A, 0x2E)  # 封面标题底线（深蓝黑）

TABLE_HEADER_BG = "404040"   # 表头深灰背景
TABLE_ROW_ALT   = "F7F7F7"  # 交替行浅灰
TABLE_BORDER    = "CCCCCC"  # 表格边线

RISK_HIGH   = RGBColor(0x00, 0x00, 0x00)  # 重大 — 加粗黑色
RISK_MED    = RGBColor(0x55, 0x55, 0x55)  # 一般 — 深灰
RISK_LOW    = RGBColor(0x88, 0x88, 0x88)  # 轻微 — 中灰


# ── 工具函数 ──────────────────────────────────────────────────────────────

def set_font(run, name="宋体", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color or BLACK
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), 'Arial')
    rFonts.set(qn('w:hAnsi'), 'Arial')
    rPr.insert(0, rFonts)


def set_para_spacing(para, before=0, after=4, line_spacing=None):
    pPr = para._p.get_or_add_pPr()
    sp  = OxmlElement('w:spacing')
    sp.set(qn('w:before'), str(before * 20))
    sp.set(qn('w:after'),  str(after  * 20))
    if line_spacing:
        sp.set(qn('w:line'),     str(int(line_spacing * 240)))
        sp.set(qn('w:lineRule'), 'auto')
    pPr.append(sp)


def set_indent(para, left_twips=360, hanging=0):
    pPr = para._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), str(left_twips))
    if hanging:
        ind.set(qn('w:hanging'), str(hanging))
    pPr.append(ind)


def set_cell_bg(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def add_table_borders(table):
    tbl   = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    bdr   = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), TABLE_BORDER)
        bdr.append(el)
    tblPr.append(bdr)


def add_hr(doc, thickness=6, color="AAAAAA"):
    """添加水平分隔线（通过段落下边框实现）"""
    para = doc.add_paragraph()
    set_para_spacing(para, before=4, after=4)
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    str(thickness))
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def risk_color(text: str) -> RGBColor:
    """根据风险程度文字返回颜色"""
    if '重大' in text:  return RISK_HIGH
    if '一般' in text:  return RISK_MED
    if '轻微' in text:  return RISK_LOW
    return BLACK


# ── Markdown 解析与渲染 ───────────────────────────────────────────────────

class LegalMarkdownToDocx:
    """
    将法律意见要点 Markdown 转换为正式 Word 文档。
    设计原则：
      - 无彩色高亮，仅使用黑/深灰/中灰
      - 每个 issue 块的字段（条款位置/问题分析等）渲染为带标签的文本行
      - 表格用简洁灰色调
      - 代码块（替换文本）用浅灰背景缩进块
    """

    def __init__(self, doc: Document):
        self.doc           = doc
        self.table_rows    = []
        self.in_table      = False
        self.code_lines    = []
        self.in_code       = False
        self.quote_lines   = []
        self.in_quote      = False

    # ── flush helpers ────────────────────────────────────────────────────

    def _flush_table(self):
        if not self.table_rows:
            return
        rows = [r for r in self.table_rows
                if not re.match(r'^[\s|:\-]+$', r.replace('|', '').strip())]
        if not rows:
            self.table_rows = []
            return

        parsed = []
        for row in rows:
            cells = [c.strip() for c in row.strip().strip('|').split('|')]
            parsed.append(cells)

        ncols = max(len(r) for r in parsed)
        tbl   = self.doc.add_table(rows=len(parsed), cols=ncols)
        tbl.style = 'Table Grid'
        add_table_borders(tbl)

        for i, row_data in enumerate(parsed):
            for j in range(min(len(row_data), ncols)):
                cell  = tbl.cell(i, j)
                ctext = row_data[j]
                p     = cell.paragraphs[0]

                # 检测风险程度文字
                is_risk = any(w in ctext for w in ['重大', '一般', '轻微'])
                r_color = risk_color(ctext) if is_risk else None

                clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', ctext)  # 去掉 **bold**
                run   = p.add_run(clean)
                set_font(run, size=9,
                         bold=(i == 0),
                         color=(RGBColor(0xFF, 0xFF, 0xFF) if i == 0 else (r_color or BLACK)))
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                if i == 0:
                    set_cell_bg(cell, TABLE_HEADER_BG)
                elif i % 2 == 0:
                    set_cell_bg(cell, TABLE_ROW_ALT)

        self.table_rows = []
        self.doc.add_paragraph()

    def _flush_code(self):
        if not self.code_lines:
            return
        para = self.doc.add_paragraph()
        set_para_spacing(para, before=2, after=6)
        set_indent(para, left_twips=480)
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F2F2F2')
        pPr.append(shd)
        for line in self.code_lines:
            run = para.add_run(line + '\n')
            run.font.name = 'Courier New'
            run.font.size = Pt(9.5)
            run.font.color.rgb = DARK_GRAY
        self.code_lines = []

    def _flush_quote(self):
        if not self.quote_lines:
            return
        text = ' '.join(self.quote_lines)
        para = self.doc.add_paragraph()
        set_para_spacing(para, before=2, after=2)
        pPr  = para._p.get_or_add_pPr()
        # 左缩进
        ind  = OxmlElement('w:ind')
        ind.set(qn('w:left'), '480')
        pPr.append(ind)
        # 细灰左边框
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'),   'single')
        left.set(qn('w:sz'),    '6')
        left.set(qn('w:space'), '4')
        left.set(qn('w:color'), 'AAAAAA')
        pBdr.append(left)
        pPr.append(pBdr)
        run = para.add_run(text)
        set_font(run, size=10, italic=True, color=DARK_GRAY)
        self.quote_lines = []

    # ── 行内格式渲染 ─────────────────────────────────────────────────────

    def _render_inline(self, para, text: str):
        """处理 **bold**、`code`、[标签] 等行内格式"""
        parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`|\[商业决策\]|\[起草技术\])', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = para.add_run(part[2:-2])
                set_font(run, bold=True)
            elif part.startswith('`') and part.endswith('`'):
                run = para.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(9.5)
                run.font.color.rgb = DARK_GRAY
            elif part in ('[商业决策]', '[起草技术]'):
                run = para.add_run(part)
                set_font(run, bold=True, color=DARK_GRAY)
            else:
                if part.strip():
                    run = para.add_run(part)
                    set_font(run)

    # ── 主处理循环 ───────────────────────────────────────────────────────

    def process_line(self, line: str):
        stripped = line.rstrip()

        # 代码块
        if stripped.startswith('```'):
            if self.in_code:
                self._flush_code()
                self.in_code = False
            else:
                if self.in_quote: self._flush_quote()
                self.in_code = True
            return
        if self.in_code:
            self.code_lines.append(stripped)
            return

        # 表格
        if '|' in stripped and stripped.strip().startswith('|'):
            if self.in_quote: self._flush_quote()
            self.table_rows.append(stripped)
            self.in_table = True
            return
        else:
            if self.in_table:
                self._flush_table()
                self.in_table = False

        # 引用
        if stripped.startswith('>'):
            content = re.sub(r'^>\s*', '', stripped)
            if content:
                self.quote_lines.append(content)
            return
        else:
            if self.quote_lines:
                self._flush_quote()

        # 空行
        if not stripped:
            if self.quote_lines: self._flush_quote()
            return

        # 水平分隔线 ---
        if re.match(r'^-{3,}$', stripped):
            if self.quote_lines: self._flush_quote()
            add_hr(self.doc)
            return

        # 一级标题 # （文档主标题）
        if stripped.startswith('# '):
            para = self.doc.add_paragraph()
            set_para_spacing(para, before=0, after=12)
            run  = para.add_run(stripped[2:].strip())
            set_font(run, size=16, bold=True, color=DARK_GRAY)
            # 底部细线
            pPr  = para._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bot  = OxmlElement('w:bottom')
            bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '4')
            bot.set(qn('w:space'), '4');    bot.set(qn('w:color'), '404040')
            pBdr.append(bot); pPr.append(pBdr)
            return

        # 二级标题 ##
        if stripped.startswith('## '):
            para = self.doc.add_paragraph()
            set_para_spacing(para, before=14, after=4)
            run  = para.add_run(stripped[3:].strip())
            set_font(run, size=13, bold=True, color=DARK_GRAY)
            return

        # 三级标题 ###
        if stripped.startswith('### '):
            para = self.doc.add_paragraph()
            set_para_spacing(para, before=10, after=4)
            run  = para.add_run(stripped[4:].strip())
            set_font(run, size=11, bold=True, color=DARK_GRAY)
            return

        # 四级标题 #### （细节标题）
        if stripped.startswith('#### '):
            para = self.doc.add_paragraph()
            set_para_spacing(para, before=8, after=2)
            run  = para.add_run(stripped[5:].strip())
            set_font(run, size=11, bold=True, color=MED_GRAY)
            return

        # 缩进块（4空格 or Tab）— 用于替换文本示例
        if stripped.startswith('    ') or stripped.startswith('\t'):
            para = self.doc.add_paragraph()
            set_para_spacing(para, before=0, after=2)
            set_indent(para, left_twips=720)
            pPr  = para._p.get_or_add_pPr()
            shd  = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'F2F2F2')
            pPr.append(shd)
            run  = para.add_run(stripped.lstrip())
            run.font.name = 'Courier New'
            run.font.size = Pt(9.5)
            run.font.color.rgb = DARK_GRAY
            return

        # 列表项 -
        if re.match(r'^[-*]\s', stripped):
            para = self.doc.add_paragraph()
            set_para_spacing(para, before=0, after=2)
            set_indent(para, left_twips=360, hanging=180)
            run  = para.add_run('•  ')
            set_font(run, bold=True, color=MED_GRAY)
            self._render_inline(para, stripped[2:].strip())
            return

        # 有序列表 1.
        if re.match(r'^\d+\.\s', stripped):
            m    = re.match(r'^(\d+\.\s)', stripped)
            para = self.doc.add_paragraph()
            set_para_spacing(para, before=0, after=2)
            set_indent(para, left_twips=360, hanging=180)
            run  = para.add_run(m.group(1))
            set_font(run, bold=True, color=DARK_GRAY)
            self._render_inline(para, stripped[m.end():].strip())
            return

        # 普通段落
        para = self.doc.add_paragraph()
        set_para_spacing(para, before=0, after=4, line_spacing=1.35)
        self._render_inline(para, stripped)

    def flush_all(self):
        if self.in_code:    self._flush_code()
        if self.table_rows: self._flush_table()
        if self.quote_lines: self._flush_quote()


# ── 文档构建 ──────────────────────────────────────────────────────────────

def build_docx(markdown_text: str, project_name: str, output_path: str):
    doc = Document()

    # 页面设置：A4，标准法律文书边距
    section = doc.sections[0]
    section.page_width   = Cm(21.0)
    section.page_height  = Cm(29.7)
    section.left_margin  = section.right_margin  = Cm(3.0)
    section.top_margin   = section.bottom_margin = Cm(2.5)

    # ── 封面 ──────────────────────────────────────────────────────────────

    # 主标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(title_para, before=40, after=4)
    run = title_para.add_run('法律审查意见要点')
    set_font(run, size=20, bold=True, color=DARK_GRAY)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(sub_para, before=0, after=20)
    run = sub_para.add_run('（草稿）')
    set_font(run, size=11, color=MED_GRAY)

    # 分隔线
    add_hr(doc, thickness=8, color="333333")

    # 项目名称和基本信息
    project_para = doc.add_paragraph()
    project_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(project_para, before=12, after=4)
    run = project_para.add_run(project_name)
    set_font(run, size=14, bold=True, color=DARK_GRAY)

    # 提取评分和评级
    score_match = re.search(
        r'综合评分[：:].*?(\d+)/100|风险评级.*?（综合评分[：:](\d+)/100）|安全评分[：:][^\d]*(\d+)',
        markdown_text
    )
    score = None
    if score_match:
        score = int(next(g for g in score_match.groups() if g))

    rating_match = re.search(r'风险评级[：:]\s*([重大中等较低高度]+风险)', markdown_text)
    rating = rating_match.group(1) if rating_match else None

    if score is not None or rating:
        score_line = " | ".join(filter(None, [
            f"综合评分：{score}/100" if score else None,
            f"风险评级：{rating}" if rating else None,
        ]))
        score_para = doc.add_paragraph()
        score_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(score_para, before=4, after=4)
        run = score_para.add_run(score_line)
        set_font(run, size=11, color=MED_GRAY)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(date_para, before=12, after=0)
    run = date_para.add_run(datetime.now().strftime("%Y年%m月%d日"))
    set_font(run, size=10, color=MED_GRAY)

    add_hr(doc, thickness=4, color="AAAAAA")

    # 免责说明
    disc_para = doc.add_paragraph()
    disc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(disc_para, before=8, after=0)
    run = disc_para.add_run('本意见由 AI 辅助生成，仅供参考，不构成正式法律意见')
    set_font(run, size=9, italic=True, color=MED_GRAY)

    doc.add_page_break()

    # ── 正文 ──────────────────────────────────────────────────────────────
    parser = LegalMarkdownToDocx(doc)
    for line in markdown_text.split('\n'):
        parser.process_line(line)
    parser.flush_all()

    # ── 页脚 ──────────────────────────────────────────────────────────────
    footer = doc.sections[0].footer
    fp     = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(
        f'法律审查意见要点（草稿）· {project_name} · '
        f'{datetime.now().strftime("%Y年%m月%d日")} · '
        '本意见不构成正式法律意见'
    )
    set_font(run, size=8, color=MED_GRAY)

    doc.save(output_path)
    print(f'✅ 报告已保存：{output_path}')


# ── CLI ──────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description='将法律审查意见 Markdown 转换为 Word 文档')
    parser.add_argument('--input',  required=True,  help='Markdown 文件路径')
    parser.add_argument('--name',   required=True,  help='项目名称（用于命名和封面）')
    parser.add_argument('--output', required=True,  help='输出目录路径')
    args = parser.parse_args()

    with open(args.input, 'r', encoding='utf-8') as f:
        md_text = f.read()

    os.makedirs(args.output, exist_ok=True)
    timestamp  = datetime.now().strftime('%Y%m%d_%H%M')
    safe_name  = re.sub(r'[\\/:*?"<>|]', '_', args.name)
    filename   = f'{safe_name}_法律审查意见_{timestamp}.docx'
    out_path   = os.path.join(args.output, filename)

    build_docx(md_text, args.name, out_path)


if __name__ == '__main__':
    main()
